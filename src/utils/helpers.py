import streamlit as st

def initialize_session_state():
    """
    Setup default Streamlit session state keys.
    """
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "llm_url" not in st.session_state:
        st.session_state.llm_url = "http://127.0.0.1:11434/v1"
    if "llm_key" not in st.session_state:
        st.session_state.llm_key = "no-key-required"
    if "llm_model" not in st.session_state:
        st.session_state.llm_model = "llama3.2"

def inject_custom_css():
    """
    Inject vibrant dark-themed styling, fonts, and container animations.
    """
    css_content = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');
        
        /* Apply Outfit font across the app */
        html, body, [class*="css"], .stMarkdown {
            font-family: 'Outfit', sans-serif !important;
        }
        
        /* Main background & moving gradient animation */
        @keyframes gradientBG {
            0% { background-position: 0% 50%; }
            50% { background-position: 100% 50%; }
            100% { background-position: 0% 50%; }
        }
        
        .stApp {
            background: linear-gradient(-45deg, #090d16, #111827, #1e1b4b, #030712) !important;
            background-size: 400% 400% !important;
            animation: gradientBG 18s ease infinite !important;
        }
        
        /* Floating background glow blobs */
        @keyframes float {
            0% { transform: translateY(0px) translateX(0px) scale(1); }
            50% { transform: translateY(-40px) translateX(30px) scale(1.1); }
            100% { transform: translateY(0px) translateX(0px) scale(1); }
        }
        @keyframes float-reverse {
            0% { transform: translateY(0px) translateX(0px) scale(1.1); }
            50% { transform: translateY(40px) translateX(-30px) scale(1); }
            100% { transform: translateY(0px) translateX(0px) scale(1.1); }
        }
        .glow-blob {
            position: fixed !important;
            border-radius: 50% !important;
            filter: blur(130px) !important;
            opacity: 0.16 !important;
            z-index: -99 !important; /* Send to absolute back */
            pointer-events: none !important;
        }
        .blob-1 {
            width: 450px !important;
            height: 450px !important;
            background: #4f46e5 !important;
            top: -150px !important;
            left: -100px !important;
            animation: float 22s infinite ease-in-out !important;
        }
        .blob-2 {
            width: 500px !important;
            height: 500px !important;
            background: #db2777 !important;
            bottom: -200px !important;
            right: -100px !important;
            animation: float-reverse 26s infinite ease-in-out !important;
        }
        
        /* Customized cards styling */
        .metric-card {
            background: rgba(30, 41, 59, 0.45);
            backdrop-filter: blur(12px);
            border-radius: 12px;
            border: 1px solid rgba(255, 255, 255, 0.08);
            padding: 20px;
            box-shadow: 0 4px 30px rgba(0, 0, 0, 0.1);
            transition: all 0.3s ease;
        }
        .metric-card:hover {
            transform: translateY(-2px);
            border-color: rgba(99, 102, 241, 0.4);
            box-shadow: 0 10px 25px rgba(99, 102, 241, 0.15);
        }
        
        /* Profile cards styling */
        .profile-container {
            display: flex;
            gap: 20px;
            margin-top: 20px;
            flex-wrap: wrap;
        }
        
        .profile-card {
            background: rgba(30, 41, 59, 0.35);
            backdrop-filter: blur(8px);
            border-radius: 16px;
            border: 1px solid rgba(255, 255, 255, 0.06);
            padding: 25px;
            flex: 1;
            min-width: 280px;
            text-align: center;
            transition: all 0.3s ease;
        }
        .profile-card:hover {
            transform: translateY(-4px);
            border-color: rgba(129, 140, 248, 0.5);
            box-shadow: 0 12px 30px rgba(129, 140, 248, 0.12);
        }
        
        .profile-avatar {
            width: 90px;
            height: 90px;
            border-radius: 50%;
            background: linear-gradient(135deg, #6366f1 0%, #a5b4fc 100%);
            margin: 0 auto 15px auto;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 32px;
            color: white;
            font-weight: 700;
            box-shadow: 0 4px 15px rgba(99, 102, 241, 0.3);
        }
        
        .profile-name {
            font-size: 20px;
            font-weight: 600;
            color: #f8fafc;
            margin: 10px 0 5px 0;
        }
        
        .profile-role {
            font-size: 14px;
            color: #818cf8;
            font-weight: 500;
            margin-bottom: 12px;
        }
        
        .profile-desc {
            font-size: 13px;
            color: #94a3b8;
            line-height: 1.5;
        }
        
        /* Titles and headers */
        h1, h2, h3 {
            color: #f8fafc !important;
            font-weight: 700 !important;
            letter-spacing: -0.025em;
        }
        
        .gradient-text {
            background: linear-gradient(135deg, #a5b4fc 0%, #818cf8 50%, #6366f1 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }
        
        /* Sidebar styling override */
        section[data-testid="stSidebar"] {
            background-color: rgb(15, 23, 42) !important;
            border-right: 1px solid rgba(255, 255, 255, 0.08);
        }
        
        /* Premium Sidebar Radio Navigation Override (Solid Bars Style) */
        [data-testid="stSidebar"] div[data-testid="stRadio"] {
            background: transparent !important;
            padding: 0 !important;
        }
        [data-testid="stSidebar"] div[data-testid="stRadio"] > div[role="radiogroup"] {
            gap: 12px !important;
        }
        [data-testid="stSidebar"] div[data-testid="stRadio"] label[data-baseweb="radio"] {
            background: rgba(15, 23, 42, 0.4) !important;
            border: 1px solid rgba(255, 255, 255, 0.04) !important;
            border-left: 4px solid transparent !important; /* Invisible indicator for inactive bars */
            border-radius: 6px !important;
            padding: 12px 18px !important;
            margin: 0 !important;
            transition: all 0.2s cubic-bezier(0.4, 0, 0.2, 1) !important;
            cursor: pointer !important;
            width: 100% !important;
        }
        [data-testid="stSidebar"] div[data-testid="stRadio"] label[data-baseweb="radio"]:hover {
            background: rgba(99, 102, 241, 0.08) !important;
            border-color: rgba(99, 102, 241, 0.2) !important;
            border-left: 4px solid rgba(99, 102, 241, 0.4) !important;
            transform: translateX(4px) !important;
        }
        [data-testid="stSidebar"] div[data-testid="stRadio"] label[data-baseweb="radio"]:has(input:checked) {
            background: rgba(99, 102, 241, 0.14) !important;
            border-color: rgba(99, 102, 241, 0.4) !important;
            border-left: 4px solid #6366f1 !important; /* Left active bar indicator */
            box-shadow: 0 4px 15px rgba(99, 102, 241, 0.08) !important;
        }
        /* Completely hide default circular radio selectors inside the sidebar */
        [data-testid="stSidebar"] div[data-testid="stRadio"] label[data-baseweb="radio"] > div:first-child {
            display: none !important;
        }
        [data-testid="stSidebar"] div[data-testid="stRadio"] label[data-baseweb="radio"] div[data-testid="stMarkdownContainer"] p {
            font-size: 14px !important;
            font-weight: 600 !important;
            color: #94a3b8 !important;
            margin: 0 !important;
        }
        [data-testid="stSidebar"] div[data-testid="stRadio"] label[data-baseweb="radio"]:has(input:checked) div[data-testid="stMarkdownContainer"] p {
            color: #ffffff !important;
            font-weight: 700 !important;
        }
        

        /* Dynamic micro-animations for interactive components */
        div.stButton > button:first-child {
            background: linear-gradient(135deg, #6366f1 0%, #4f46e5 100%) !important;
            color: white !important;
            border: None !important;
            border-radius: 8px !important;
            font-weight: 600 !important;
            transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        }
        div.stButton > button:first-child:hover {
            transform: scale(1.02) !important;
            box-shadow: 0 0 15px rgba(99, 102, 241, 0.4) !important;
        }
        div.stButton > button:first-child:disabled {
            background: linear-gradient(135deg, #10b981 0%, #059669 100%) !important;
            color: white !important;
            opacity: 1.0 !important;
            border: None !important;
            box-shadow: 0 0 15px rgba(16, 185, 129, 0.45) !important;
            cursor: not-allowed !important;
            transform: none !important;
        }
        
        /* Diagnostic card styling */
        .diagnostic-card {
            background: rgba(15, 23, 42, 0.6);
            border-radius: 8px;
            border: 1px solid rgba(255, 255, 255, 0.05);
            padding: 15px;
            margin-top: 15px;
        }
        
        /* Custom Scrollbar Styling for all pages and modals */
        ::-webkit-scrollbar {
            width: 10px !important;
            height: 10px !important;
        }
        ::-webkit-scrollbar-track {
            background: rgba(15, 23, 42, 0.6) !important;
            border-radius: 6px !important;
        }
        ::-webkit-scrollbar-thumb {
            background: #6366f1 !important; /* Beautiful Indigo Scrollbar */
            border-radius: 6px !important;
            border: 2px solid rgba(15, 23, 42, 0.6) !important;
        }
        ::-webkit-scrollbar-thumb:hover {
            background: #4f46e5 !important;
        }
        
        /* Make Streamlit modals (like screen recording preview) scrollbar highly visible */
        [data-testid="stModal"] {
            overflow-y: auto !important;
        }
        [data-testid="stModal"] ::-webkit-scrollbar-thumb {
            background: #f59e0b !important; /* Bright orange/yellow scrollbar specifically inside the modal to draw attention! */
            border: 1px solid rgba(15, 23, 42, 0.8) !important;
        }
    </style>
    """
    st.markdown(css_content, unsafe_allow_html=True)
    
    # Inject background floating glow blobs
    st.markdown(
        """
        <div class="glow-blob blob-1"></div>
        <div class="glow-blob blob-2"></div>
        """,
        unsafe_allow_html=True
    )

def export_chat_to_docx(chat_history, ticker, prediction_context=""):
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import io
    from datetime import datetime
    import re
    
    # Helper to clean text of emojis and markdown hashes to prevent square box glyphs
    def clean_text(text):
        cleaned = "".join(c for c in text if 31 < ord(c) < 127 or c in ('\n', '\r', '\t'))
        cleaned = cleaned.replace("### ", "").replace("## ", "").replace("# ", "")
        return cleaned.strip()
        
    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)
        
    def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(m)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
        
    def set_cell_left_border(cell, color_hex, size="24"):
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), size)
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), color_hex)
        tcBorders.append(left)
        
        for b_type in ['top', 'bottom', 'right']:
            b = OxmlElement(f'w:{b_type}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
            
        tcPr.append(tcBorders)

    def add_markdown_line(container, line_text):
        cleaned_line = clean_text(line_text)
        if not cleaned_line:
            return
        p = container.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        
        # Split by bold format **
        parts = re.split(r'(\*\*.*?\*\*)', cleaned_line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Check for inline code `
                subparts = re.split(r'(`.*?`)', part)
                for subpart in subparts:
                    if subpart.startswith('`') and subpart.endswith('`'):
                        run = p.add_run(subpart[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9.5)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(79, 70, 229) # Indigo
                    else:
                        p.add_run(subpart)

    def create_code_card(container, code_text):
        table = container.add_table(rows=1, cols=1)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        set_cell_left_border(cell, "94A3B8", size="24")
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(code_text.strip())
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(15, 23, 42)

    def render_formatted_text(container, text):
        parts = re.split(r'(```[\s\S]*?```)', text)
        for part in parts:
            if part.startswith('```') and part.endswith('```'):
                # Extract code contents
                code_lines = part.split('\n')
                code_content = '\n'.join(code_lines[1:-1])
                create_code_card(container, code_content)
            else:
                lines = part.split('\n')
                for line in lines:
                    add_markdown_line(container, line)

    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Base styling
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10.5)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AI Quant-Coder Chat & Strategy Report")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(79, 70, 229)
    
    # Subtitle Metadata
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Target Asset: {ticker}  |  Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    sub_run.font.size = Pt(9.5)
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    
    # Market Context Card
    if prediction_context:
        doc.add_heading("Live Market Context", level=1)
        doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(79, 70, 229)
        doc.paragraphs[-1].runs[0].font.size = Pt(13)
        
        table_ctx = doc.add_table(rows=1, cols=1)
        cell_ctx = table_ctx.cell(0, 0)
        set_cell_background(cell_ctx, "F8FAFC")
        set_cell_margins(cell_ctx, top=140, bottom=140, left=180, right=180)
        set_cell_left_border(cell_ctx, "6366F1", size="24")
        
        ctx_lines = prediction_context.split('\n')
        for line in ctx_lines:
            add_markdown_line(cell_ctx, line)
            
    doc.add_paragraph() # Spacing
    
    # Conversation details
    doc.add_heading("Conversation & Strategy Details", level=1)
    doc.paragraphs[-1].runs[0].font.color.rgb = RGBColor(79, 70, 229)
    doc.paragraphs[-1].runs[0].font.size = Pt(13)
    
    for msg in chat_history:
        is_user = (msg["role"] == "user")
        role_label = "User Request" if is_user else "Quant-Coder Strategy & Code"
        
        # Message bubble card
        table_msg = doc.add_table(rows=1, cols=1)
        cell_msg = table_msg.cell(0, 0)
        set_cell_margins(cell_msg, top=140, bottom=140, left=200, right=200)
        
        if is_user:
            set_cell_background(cell_msg, "FFF7ED") # light orange
            set_cell_left_border(cell_msg, "F97316", size="24")
            header_color = RGBColor(234, 88, 12)
        else:
            set_cell_background(cell_msg, "EEF2FF") # light blue/indigo
            set_cell_left_border(cell_msg, "4F46E5", size="24")
            header_color = RGBColor(55, 48, 163)
            
        p_hdr = cell_msg.paragraphs[0]
        p_hdr.paragraph_format.space_after = Pt(6)
        run_hdr = p_hdr.add_run(role_label)
        run_hdr.font.bold = True
        run_hdr.font.size = Pt(11)
        run_hdr.font.color.rgb = header_color
        
        # Render contents (which parses nested markdown and code blocks!)
        render_formatted_text(cell_msg, msg["content"])
        
        # If assistant has retrieved templates, append them inside the message card!
        if not is_user and msg.get("sources"):
            p_src_title = cell_msg.add_paragraph()
            p_src_title.paragraph_format.space_before = Pt(8)
            p_src_title.paragraph_format.space_after = Pt(4)
            run_src_title = p_src_title.add_run("Retrieved Code Templates Used:")
            run_src_title.bold = True
            run_src_title.font.size = Pt(10)
            run_src_title.font.color.rgb = RGBColor(79, 70, 229)
            
            for s in msg["sources"]:
                p_src_info = cell_msg.add_paragraph()
                p_src_info.paragraph_format.space_after = Pt(2)
                run_src_info = p_src_info.add_run(f"- File: {s['file']} (Relevance: {s['score']:.2f})")
                run_src_info.italic = True
                run_src_info.font.size = Pt(9)
                create_code_card(cell_msg, s["code"])
                
        doc.add_paragraph() # spacing between messages
        
    # Footer disclaimer
    doc.add_paragraph("\n---\nReport compiled automatically by AI Quant-Coder. AMD Developer Hackathon ACT II.")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def export_chat_to_markdown(chat_history, ticker, prediction_context=""):
    import io
    from datetime import datetime
    md = f"# AI Quant-Coder Chat & Strategy Report\n\n"
    md += f"**Target Asset:** {ticker}  \n"
    md += f"**Generated on:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \n\n"
    
    if prediction_context:
        md += f"## 📊 Live Market Context\n"
        md += f"```\n{prediction_context.strip()}\n```\n\n"
        
    md += f"## 💬 Conversation Details\n\n"
    for msg in chat_history:
        role_name = "User" if msg["role"] == "user" else "AI Assistant"
        md += f"### 👤 {role_name}\n"
        md += f"{msg['content']}\n\n"
        
        if msg.get("sources"):
            md += f"#### 🔍 Retrieved Code Templates Used\n"
            for s in msg["sources"]:
                md += f"- **File:** `{s['file']}` (Score: {s['score']:.2f})\n"
                md += f"```python\n{s['code']}\n```\n\n"
                
    file_stream = io.BytesIO(md.encode('utf-8'))
    file_stream.seek(0)
    return file_stream
